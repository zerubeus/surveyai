"""
Chisquare — PDF/Word Questionnaire Parser

Uses text extraction (pypdf for PDF, python-docx for DOCX) followed by
Gemini AI to identify question structure from unstructured document text.

Returns the same ParsedInstrument format as the XLSForm parser,
but with lower confidence since it relies on AI extraction.
"""

from __future__ import annotations

import io
import json

import structlog

from parsers.xlsform_parser import ParsedInstrument, ParsedQuestion
from services.ai_service import generate

logger = structlog.get_logger()

# Gemini prompt for extracting questionnaire structure from text
EXTRACTION_PROMPT = """You are a survey research expert. Extract the questionnaire structure from the following document text.

Return a JSON object with this exact schema:
{{
  "title": "form title or best guess",
  "form_id": "form identifier if found, else empty string",
  "version": "version if found, else empty string",
  "default_language": "primary language of the document",
  "languages": ["list of languages detected"],
  "questions": [
    {{
      "name": "short variable name (e.g. q1, q2a)",
      "label": {{"default": "full question text"}},
      "type": "one of: select_one, select_multiple, integer, decimal, text, date, note",
      "choices": [
        {{"name": "1", "label": {{"default": "choice text"}}}}
      ],
      "relevant": "skip logic expression if mentioned, else null",
      "constraint": "validation rule if mentioned, else null",
      "required": false,
      "hint": {{}},
      "group_path": ["section name if applicable"]
    }}
  ]
}}

Rules:
- Identify all questions including sub-questions
- For multiple choice / rating scale questions use "select_one"
- For "check all that apply" questions use "select_multiple"
- For open-ended questions use "text"
- For numeric questions use "integer" or "decimal"
- Generate short variable names (q1, q2, q2a, q3) if not explicitly labeled
- Detect section/group structure from headings
- If response options are provided, include them in choices
- Include ALL questions, even instructions/notes (type="note")

Document text:
{text}
"""


def _extract_pdf_text(file_bytes: bytes) -> str:
    """Extract text from a PDF file using pypdf."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    pages: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def _extract_docx_text(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def _normalize_question_type(raw_type: str) -> str:
    """Map AI-returned type strings to normalized question_type."""
    mapping: dict[str, str] = {
        "select_one": "single_select",
        "select_multiple": "multi_select",
        "integer": "numeric",
        "decimal": "numeric",
        "text": "text",
        "date": "date",
        "note": "note",
        "calculate": "calculated",
        "geopoint": "gps",
    }
    return mapping.get(raw_type.lower().split()[0], "text")


def parse_document(file_bytes: bytes, mime_type: str) -> ParsedInstrument:
    """
    Parse a PDF or DOCX questionnaire file using text extraction + Gemini AI.

    Args:
        file_bytes: Raw bytes of the file
        mime_type: MIME type (application/pdf or application/vnd.openxmlformats-officedocument.wordprocessingml.document)

    Returns:
        ParsedInstrument with AI-extracted question structure

    Raises:
        ValueError: If text extraction or AI parsing fails
    """
    # Extract text based on file type
    if mime_type == "application/pdf":
        text = _extract_pdf_text(file_bytes)
    elif mime_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        text = _extract_docx_text(file_bytes)
    else:
        raise ValueError(f"Unsupported MIME type for document parsing: {mime_type}")

    if not text.strip():
        raise ValueError("No text could be extracted from the document")

    # Truncate very long documents to avoid token limits
    max_chars = 30000
    if len(text) > max_chars:
        text = text[:max_chars] + "\n\n[... document truncated ...]"
        logger.warning("document_truncated", original_length=len(text), max_chars=max_chars)

    # Send to Gemini for structure extraction
    prompt = EXTRACTION_PROMPT.format(text=text)
    ai_result = generate(prompt)

    # Parse AI response into ParsedInstrument
    questions: list[ParsedQuestion] = []
    skip_logic_count = 0

    for q_data in ai_result.get("questions", []):
        raw_type = q_data.get("type", "text")
        question_type = _normalize_question_type(raw_type)

        # Parse label
        label_raw = q_data.get("label", {})
        if isinstance(label_raw, str):
            label = {"default": label_raw}
        elif isinstance(label_raw, dict):
            label = {k: str(v) for k, v in label_raw.items()}
        else:
            label = {}

        # Parse hint
        hint_raw = q_data.get("hint", {})
        if isinstance(hint_raw, str):
            hint = {"default": hint_raw} if hint_raw else {}
        elif isinstance(hint_raw, dict):
            hint = {k: str(v) for k, v in hint_raw.items()}
        else:
            hint = {}

        # Parse choices
        choices: list[dict[str, str | dict[str, str]]] = []
        for c in q_data.get("choices", []):
            if isinstance(c, dict):
                c_label = c.get("label", {})
                if isinstance(c_label, str):
                    c_label = {"default": c_label}
                choices.append({
                    "name": str(c.get("name", "")),
                    "label": c_label,
                })

        relevant = q_data.get("relevant") or None
        if relevant:
            skip_logic_count += 1

        # Parse group_path
        group_path_raw = q_data.get("group_path", [])
        group_path = [str(g) for g in group_path_raw] if isinstance(group_path_raw, list) else []

        questions.append(ParsedQuestion(
            name=str(q_data.get("name", f"q{len(questions) + 1}")),
            label=label,
            type=raw_type,
            question_type=question_type,
            choices=choices,
            relevant=relevant,
            constraint=q_data.get("constraint") or None,
            required=bool(q_data.get("required", False)),
            hint=hint,
            appearance=None,
            group_path=group_path,
        ))

    languages = ai_result.get("languages", [])
    if isinstance(languages, list):
        languages = [str(lang) for lang in languages]
    else:
        languages = []

    default_language = str(ai_result.get("default_language", "")) or (languages[0] if languages else "English")

    logger.info(
        "document_parsed_with_ai",
        question_count=len(questions),
        skip_logic_count=skip_logic_count,
        languages=languages,
    )

    return ParsedInstrument(
        title=str(ai_result.get("title", "Untitled Form")),
        form_id=str(ai_result.get("form_id", "")),
        version=str(ai_result.get("version", "")),
        default_language=default_language,
        languages=languages,
        questions=questions,
        skip_logic_count=skip_logic_count,
        has_skip_logic=skip_logic_count > 0,
    )
