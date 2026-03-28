"""
SurveyAI Analyst — Report Export Service

Exports reports to DOCX and PDF formats.
DOCX uses python-docx, PDF uses ReportLab (pure Python, no system deps).

Storage: files uploaded to Supabase Storage 'reports' bucket.
Signed URLs (1-hour expiry) stored in report_exports table.
"""

from __future__ import annotations

import io
import json
import re
from datetime import datetime, timezone, timedelta
from typing import Any

import structlog
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from db import SupabaseDB

logger = structlog.get_logger()

# ReportLab for PDF — pure Python, no system library dependencies


def export_report(db: SupabaseDB, task_id: str, payload: dict[str, Any]) -> None:
    """
    Export report to DOCX and PDF.

    Payload: {"report_id": str, "formats": ["docx", "pdf"]}
    """
    report_id: str = payload["report_id"]
    formats: list[str] = payload.get("formats", ["docx", "pdf"])
    created_by: str = payload.get("created_by", "worker")

    logger.info("export_start", report_id=report_id, formats=formats)

    # Step 1: Load report + sections
    db.update_task_progress(task_id, 5, "Loading report data...")
    reports = db.select("reports", filters={"id": report_id})
    if not reports:
        raise ValueError(f"Report {report_id} not found")
    report = reports[0]

    project = db.get_project(report["project_id"])
    project_name = project["name"] if project else "Report"

    sections = db.select("report_sections", filters={"report_id": report_id})
    sections.sort(key=lambda s: s.get("sort_order", 0))

    # Step 2: Load chart images
    db.update_task_progress(task_id, 10, "Loading chart images...")
    chart_images = _load_chart_images(db, sections)

    # Step 3: Generate DOCX
    export_results: list[dict[str, Any]] = []

    if "docx" in formats:
        db.update_task_progress(task_id, 20, "Generating DOCX...")
        docx_bytes = _generate_docx(report, project_name, sections, chart_images)
        docx_path = f"{created_by}/{report['project_id']}/report_{report_id}.docx"

        db.upload_file("reports", docx_path, docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        signed_url = db.get_signed_url("reports", docx_path, expires_in=3600)
        expires_at = (datetime.now(timezone.utc) + timedelta(hours=1)).isoformat()

        export_record = db.insert("report_exports", {
            "report_id": report_id,
            "format": "docx",
            "file_path": docx_path,
            "file_size_bytes": len(docx_bytes),
            "generated_by": created_by,
            "expires_at": expires_at,
        })
        export_results.append({
            "format": "docx",
            "export_id": export_record["id"],
            "signed_url": signed_url,
            "file_size_bytes": len(docx_bytes),
        })
        logger.info("docx_exported", path=docx_path, size=len(docx_bytes))

    # Step 4: Generate PDF
    if "pdf" in formats:
        db.update_task_progress(task_id, 60, "Generating PDF...")
        pdf_bytes = _generate_pdf(report, project_name, sections, chart_images)
        pdf_path = f"{created_by}/{report['project_id']}/report_{report_id}.pdf"

        db.upload_file("reports", pdf_path, pdf_bytes, "application/pdf")
        signed_url = db.get_signed_url("reports", pdf_path, expires_in=3600)
        expires_at = (datetime.now(timezone.utc) + timedelta(hours=1)).isoformat()

        export_record = db.insert("report_exports", {
            "report_id": report_id,
            "format": "pdf",
            "file_path": pdf_path,
            "file_size_bytes": len(pdf_bytes),
            "generated_by": created_by,
            "expires_at": expires_at,
        })
        export_results.append({
            "format": "pdf",
            "export_id": export_record["id"],
            "signed_url": signed_url,
            "file_size_bytes": len(pdf_bytes),
        })
        logger.info("pdf_exported", path=pdf_path, size=len(pdf_bytes))

    # Step 5: Update report status
    db.update_task_progress(task_id, 95, "Finalizing export...")
    db.update("reports", {"status": "exported"}, {"id": report_id})

    db.complete_task(task_id, {
        "message": "Report exported successfully",
        "report_id": report_id,
        "exports": export_results,
    })


def _load_chart_images(
    db: SupabaseDB,
    sections: list[dict[str, Any]],
) -> dict[str, bytes]:
    """Load chart images from storage for embedding in exports."""
    images: dict[str, bytes] = {}

    for section in sections:
        linked = section.get("linked_charts")
        if not linked:
            continue

        chart_ids = linked if isinstance(linked, list) else json.loads(linked)
        for chart_id in chart_ids:
            if chart_id in images:
                continue
            try:
                charts = db.select("charts", filters={"id": chart_id})
                if charts and charts[0].get("file_path"):
                    image_data = db.download_file("charts", charts[0]["file_path"])
                    images[chart_id] = image_data
            except Exception as e:
                logger.warning("chart_image_load_failed", chart_id=chart_id, error=str(e))

    return images


def _generate_docx(
    report: dict[str, Any],
    project_name: str,
    sections: list[dict[str, Any]],
    chart_images: dict[str, bytes],
) -> bytes:
    """Generate DOCX using python-docx."""
    doc = Document()

    # Configure default style
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.space_after = Pt(6)
    run = title_para.add_run(project_name)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 51, 102)

    # Template info
    template = report.get("template", "donor")
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_para.add_run(f"Report Template: {template.title()}")
    meta_run.font.size = Pt(12)
    meta_run.font.color.rgb = RGBColor(100, 100, 100)

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(datetime.now(timezone.utc).strftime("%B %d, %Y"))
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # Sections
    for section in sections:
        doc.add_heading(section["title"], level=1)

        content = section.get("content") or ""

        # Style [EXPERT INPUT:] placeholders in red italic
        if "[EXPERT INPUT:" in content:
            _add_content_with_placeholders(doc, content)
        else:
            _add_markdown_content(doc, content)

        # Embed charts
        linked = section.get("linked_charts")
        if linked:
            chart_ids = linked if isinstance(linked, list) else json.loads(linked)
            for chart_id in chart_ids:
                if chart_id in chart_images:
                    try:
                        image_stream = io.BytesIO(chart_images[chart_id])
                        doc.add_picture(image_stream, width=Inches(5.5))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        logger.warning("chart_embed_failed", chart_id=chart_id, error=str(e))

        # Add spacing between sections
        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _add_content_with_placeholders(doc: Document, content: str) -> None:
    """Add content with [EXPERT INPUT:] styled in red italic."""
    # Split on placeholder patterns
    parts = re.split(r"(\[EXPERT INPUT:[^\]]*\])", content)
    for part in parts:
        if part.startswith("[EXPERT INPUT:"):
            para = doc.add_paragraph()
            run = para.add_run(part)
            run.italic = True
            run.font.color.rgb = RGBColor(204, 0, 0)
            run.font.size = Pt(11)
        elif part.strip():
            _add_markdown_content(doc, part)


def _add_markdown_content(doc: Document, content: str) -> None:
    """Convert markdown content to docx paragraphs."""
    lines = content.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Handle markdown headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            para = doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            para = doc.add_paragraph(text, style="List Number")
        else:
            para = doc.add_paragraph()
            # Handle bold and italic within text
            _add_formatted_runs(para, stripped)


def _add_formatted_runs(para, text: str) -> None:
    """Add formatted runs handling **bold** and *italic* markdown."""
    # Simple bold/italic parsing
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            para.add_run(part)


def _generate_pdf(
    report: dict[str, Any],
    project_name: str,
    sections: list[dict[str, Any]],
    chart_images: dict[str, bytes],
) -> bytes:
    """Generate PDF using ReportLab — pure Python, no system library dependencies."""
    from io import BytesIO
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, HRFlowable
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("title", parent=styles["Title"], fontSize=24, spaceAfter=12, textColor=colors.HexColor("#003366"))
    h1_style = ParagraphStyle("h1", parent=styles["Heading1"], fontSize=16, textColor=colors.HexColor("#003366"), spaceAfter=8)
    h2_style = ParagraphStyle("h2", parent=styles["Heading2"], fontSize=13, textColor=colors.HexColor("#004488"), spaceAfter=6)
    body_style = ParagraphStyle("body", parent=styles["Normal"], fontSize=11, leading=16, spaceAfter=8)
    placeholder_style = ParagraphStyle("placeholder", parent=styles["Normal"], fontSize=11, textColor=colors.red, fontName="Helvetica-Oblique")
    review_style = ParagraphStyle("review", parent=styles["Normal"], fontSize=10, textColor=colors.HexColor("#795548"), fontName="Helvetica-Oblique")

    story = []

    # Title page
    story.append(Spacer(1, 6 * cm))
    story.append(Paragraph(project_name, title_style))
    story.append(Paragraph(f"Report Template: {report.get('template', 'donor').title()}", body_style))
    story.append(Spacer(1, 1 * cm))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#003366")))
    story.append(Spacer(1, 2 * cm))

    for section in sorted(sections, key=lambda s: s.get("section_order", 0)):
        content = section.get("content", "")
        if not content:
            continue

        title = section.get("title", section.get("section_type", "").replace("_", " ").title())
        confidence = section.get("confidence_level", "high")
        chart_path = section.get("chart_path")

        story.append(Paragraph(title, h1_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cccccc")))
        story.append(Spacer(1, 0.3 * cm))

        if confidence == "medium":
            story.append(Paragraph("⚠ This section may need expert review.", review_style))
            story.append(Spacer(1, 0.2 * cm))

        # Split content into paragraphs, handle placeholders
        for para in content.split("\n\n"):
            para = para.strip()
            if not para:
                continue
            if "[EXPERT INPUT:" in para:
                story.append(Paragraph(para, placeholder_style))
            else:
                story.append(Paragraph(para.replace("\n", "<br/>"), body_style))

        # Embed chart if available
        if chart_path and chart_path in chart_images:
            img_data = BytesIO(chart_images[chart_path])
            img = RLImage(img_data, width=14 * cm, height=9 * cm)
            story.append(Spacer(1, 0.5 * cm))
            story.append(img)

        story.append(Spacer(1, 1 * cm))

    doc.build(story)
    return buffer.getvalue()


def _build_html(
    project_name: str,
    template: str,
    sections: list[dict[str, Any]],
    chart_images: dict[str, bytes],
) -> str:
    """Build styled HTML for PDF generation."""
    import base64

    css = """
    body {
        font-family: 'Helvetica Neue', Arial, sans-serif;
        font-size: 11pt;
        line-height: 1.6;
        color: #333;
        margin: 2cm 2.5cm;
    }
    .title-page {
        text-align: center;
        padding-top: 30%;
        page-break-after: always;
    }
    .title-page h1 {
        font-size: 28pt;
        color: #003366;
        margin-bottom: 0.5em;
    }
    .title-page .meta {
        font-size: 12pt;
        color: #666;
    }
    h1 { font-size: 18pt; color: #003366; border-bottom: 2px solid #003366; padding-bottom: 4px; margin-top: 1.5em; }
    h2 { font-size: 14pt; color: #004488; margin-top: 1.2em; }
    h3 { font-size: 12pt; color: #555; margin-top: 1em; }
    .placeholder {
        color: #cc0000;
        font-style: italic;
        background-color: #fff5f5;
        padding: 8px 12px;
        border-left: 3px solid #cc0000;
        margin: 1em 0;
    }
    .review-banner {
        background-color: #fff8e1;
        border-left: 3px solid #f9a825;
        padding: 8px 12px;
        margin: 1em 0;
        font-size: 10pt;
        color: #795548;
    }
    .chart-container {
        text-align: center;
        margin: 1.5em 0;
    }
    .chart-container img {
        max-width: 100%;
        height: auto;
    }
    table { border-collapse: collapse; width: 100%; margin: 1em 0; }
    th, td { border: 1px solid #ddd; padding: 6px 8px; text-align: left; }
    th { background-color: #f5f5f5; }
    """

    body_parts = []

    # Title page
    date_str = datetime.now(timezone.utc).strftime("%B %d, %Y")
    body_parts.append(f"""
    <div class="title-page">
        <h1>{_html_escape(project_name)}</h1>
        <div class="meta">Report Template: {template.title()}</div>
        <div class="meta">{date_str}</div>
    </div>
    """)

    # Sections
    for section in sections:
        content = section.get("content") or ""
        confidence = section.get("confidence")
        has_placeholders = section.get("has_placeholders", False)

        body_parts.append(f"<h1>{_html_escape(section['title'])}</h1>")

        if confidence == "medium":
            body_parts.append('<div class="review-banner">⚠ This section was AI-generated and needs review.</div>')

        if has_placeholders or "[EXPERT INPUT:" in content:
            # Style placeholder content
            placeholder_parts = re.split(r"(\[EXPERT INPUT:[^\]]*\])", content)
            for part in placeholder_parts:
                if part.startswith("[EXPERT INPUT:"):
                    body_parts.append(f'<div class="placeholder">{_html_escape(part)}</div>')
                elif part.strip():
                    body_parts.append(_markdown_to_html(part))
        else:
            body_parts.append(_markdown_to_html(content))

        # Embed charts
        linked = section.get("linked_charts")
        if linked:
            chart_ids = linked if isinstance(linked, list) else json.loads(linked)
            for chart_id in chart_ids:
                if chart_id in chart_images:
                    b64 = base64.b64encode(chart_images[chart_id]).decode()
                    body_parts.append(f"""
                    <div class="chart-container">
                        <img src="data:image/png;base64,{b64}" alt="Chart" />
                    </div>
                    """)

    body_html = "\n".join(body_parts)

    return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>{css}</style>
</head>
<body>
    {body_html}
</body>
</html>"""


def _html_escape(text: str) -> str:
    """Escape HTML special characters."""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _markdown_to_html(content: str) -> str:
    """Simple markdown to HTML conversion."""
    lines = content.split("\n")
    html_parts = []
    in_list = False
    list_type = ""

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if in_list:
                html_parts.append(f"</{list_type}>")
                in_list = False
            html_parts.append("")
            continue

        # Headings
        if stripped.startswith("### "):
            if in_list:
                html_parts.append(f"</{list_type}>")
                in_list = False
            html_parts.append(f"<h3>{_html_escape(stripped[4:])}</h3>")
        elif stripped.startswith("## "):
            if in_list:
                html_parts.append(f"</{list_type}>")
                in_list = False
            html_parts.append(f"<h2>{_html_escape(stripped[3:])}</h2>")
        elif stripped.startswith("# "):
            if in_list:
                html_parts.append(f"</{list_type}>")
                in_list = False
            html_parts.append(f"<h2>{_html_escape(stripped[2:])}</h2>")  # h2 since h1 is section title
        elif stripped.startswith("- ") or stripped.startswith("* "):
            if not in_list or list_type != "ul":
                if in_list:
                    html_parts.append(f"</{list_type}>")
                html_parts.append("<ul>")
                in_list = True
                list_type = "ul"
            html_parts.append(f"<li>{_format_inline(stripped[2:])}</li>")
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            if not in_list or list_type != "ol":
                if in_list:
                    html_parts.append(f"</{list_type}>")
                html_parts.append("<ol>")
                in_list = True
                list_type = "ol"
            html_parts.append(f"<li>{_format_inline(text)}</li>")
        else:
            if in_list:
                html_parts.append(f"</{list_type}>")
                in_list = False
            html_parts.append(f"<p>{_format_inline(stripped)}</p>")

    if in_list:
        html_parts.append(f"</{list_type}>")

    return "\n".join(html_parts)


def _format_inline(text: str) -> str:
    """Format inline markdown (bold, italic) to HTML."""
    escaped = _html_escape(text)
    # Bold
    escaped = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", escaped)
    # Italic
    escaped = re.sub(r"\*([^*]+)\*", r"<em>\1</em>", escaped)
    return escaped
